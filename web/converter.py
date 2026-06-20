import subprocess
import tempfile
import os
import re
import shutil
import io
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


# ─── PDF Conversion (Java JAR) ──────────────────────────────────────────────

JAR_PATH = Path(__file__).parent.parent / "java" / "opendataloader-pdf-cli" / "target" / "opendataloader-pdf-cli-0.0.0.jar"


def _find_jar() -> Path:
    if JAR_PATH.exists():
        return JAR_PATH
    target_dir = Path(__file__).parent.parent / "java" / "opendataloader-pdf-cli" / "target"
    if target_dir.exists():
        for f in target_dir.glob("opendataloader-pdf-cli-*.jar"):
            if "original-" not in f.name and "sources" not in f.name and "javadoc" not in f.name:
                return f
    raise FileNotFoundError(
        f"Java JAR not found at {JAR_PATH}. Build it first: cd java && mvn package -DskipTests"
    )


def _find_java() -> str:
    java_path = shutil.which("java")
    if java_path:
        return java_path
    for candidate in [
        r"C:\Program Files\Eclipse Adoptium\jdk-25.0.3.9-hotspot\bin\java.exe",
        r"C:\Program Files\Eclipse Adoptium\jdk-21\bin\java.exe",
        r"C:\Program Files\Java\jdk-21\bin\java.exe",
        r"C:\Program Files\Java\jdk-17\bin\java.exe",
    ]:
        if os.path.exists(candidate):
            return candidate
    raise FileNotFoundError("Java not found on PATH or in common locations")


def run_java_jar(pdf_path: str, output_dir: str, fmt: str) -> dict[str, Path]:
    jar = _find_jar()
    java = _find_java()
    cmd = [
        java, "-Djava.awt.headless=true", "-Dapple.awt.UIElement=true",
        "-jar", str(jar),
        "-f", fmt,
        "-o", output_dir,
        str(pdf_path),
    ]
    result = subprocess.run(cmd, capture_output=True, text=True, timeout=300)
    if result.returncode != 0:
        raise RuntimeError(f"Java JAR failed (code {result.returncode}):\n{result.stderr}\n{result.stdout}")

    stem = Path(pdf_path).stem
    outputs = {}
    for ext in ("md", "json", "html", "txt"):
        p = Path(output_dir) / f"{stem}.{ext}"
        if p.exists():
            outputs[ext] = p
    return outputs


def pdf_to_markdown(pdf_path: str, output_dir: str) -> Path:
    outputs = run_java_jar(pdf_path, output_dir, "markdown")
    if "md" not in outputs:
        raise RuntimeError("Markdown output not generated")
    return outputs["md"]


def pdf_to_json(pdf_path: str, output_dir: str) -> Path:
    outputs = run_java_jar(pdf_path, output_dir, "json")
    if "json" not in outputs:
        raise RuntimeError("JSON output not generated")
    return outputs["json"]


def _escape_docx_xml(text: str) -> str:
    return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def _add_heading(doc: Document, line: str):
    level = len(line) - len(line.lstrip("#"))
    text = line.lstrip("# ").strip()
    if level > 9:
        level = 9
    doc.add_heading(text, level=level)


def _parse_md_table(lines: list[str]) -> tuple[list[str], list[list[str]]]:
    headers = [c.strip() for c in lines[0].strip("|").split("|")]
    rows = []
    for line in lines[2:]:
        cells = [c.strip() for c in line.strip("|").split("|")]
        rows.append(cells)
    return headers, rows


def _add_table(doc: Document, headers: list[str], rows: list[list[str]]):
    num_cols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=num_cols, style="Table Grid")
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
    for r_idx, row in enumerate(rows):
        for c_idx, cell_text in enumerate(row):
            if c_idx < num_cols:
                table.rows[r_idx + 1].cells[c_idx].text = cell_text
    doc.add_paragraph()


def _add_list_item(doc: Document, line: str):
    indent = len(line) - len(line.lstrip())
    text = re.sub(r"^[\s]*[-*+]\s+", "", line)
    text = re.sub(r"^[\s]*\d+\.\s+", "", text)
    text = text.strip()
    p = doc.add_paragraph(text, style="List Bullet")
    if indent >= 2:
        p.paragraph_format.left_indent = Inches(0.5 * (indent // 2))


def markdown_to_docx(md_path: str, output_path: str) -> Path:
    doc = Document()

    style = doc.styles["Normal"]
    style.font.size = Pt(11)
    style.font.name = "Calibri"

    with open(md_path, "r", encoding="utf-8") as f:
        lines = f.readlines()

    i = 0
    in_code_block = False
    code_lines = []

    while i < len(lines):
        line = lines[i].rstrip("\n")

        if line.strip().startswith("```"):
            if in_code_block:
                code_text = "\n".join(code_lines)
                p = doc.add_paragraph()
                run = p.add_run(code_text)
                run.font.name = "Courier New"
                run.font.size = Pt(9)
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        if re.match(r"^#{1,6}\s", stripped):
            _add_heading(doc, stripped)
            i += 1
            continue

        if "|" in stripped and i + 1 < len(lines) and re.match(r"^\s*\|[\s:|-]+\|\s*$", lines[i + 1]):
            table_lines = [stripped]
            j = i + 2
            while j < len(lines) and "|" in lines[j].strip():
                table_lines.append(lines[j].strip())
                j += 1
            headers, rows = _parse_md_table(table_lines)
            _add_table(doc, headers, rows)
            i = j
            continue

        if re.match(r"^[\s]*[-*+]\s", line) or re.match(r"^[\s]*\d+\.\s", line):
            _add_list_item(doc, line)
            i += 1
            continue

        if stripped.startswith("> "):
            text = stripped[2:]
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.italic = True
            run.font.color.rgb = RGBColor(100, 100, 100)
            i += 1
            continue

        if stripped.startswith("---") or stripped.startswith("***"):
            doc.add_paragraph("_" * 40)
            i += 1
            continue

        doc.add_paragraph(stripped)
        i += 1

    doc.save(output_path)
    return Path(output_path)


def pdf_to_docx(pdf_path: str, output_dir: str) -> Path:
    md_path = pdf_to_markdown(pdf_path, output_dir)
    docx_path = Path(output_dir) / (Path(pdf_path).stem + ".docx")
    return markdown_to_docx(str(md_path), str(docx_path))


FORMAT_HANDLERS = {
    "md": pdf_to_markdown,
    "json": pdf_to_json,
    "docx": pdf_to_docx,
}

FORMAT_MIME = {
    "md": "text/markdown",
    "json": "application/json",
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}

FORMAT_EXT = {
    "md": ".md",
    "json": ".json",
    "docx": ".docx",
}


# ─── Image Conversion ────────────────────────────────────────────────────────

from PIL import Image

IMG_FORMATS = {
    "png": "PNG",
    "jpg": "JPEG",
    "jpeg": "JPEG",
    "webp": "WEBP",
    "bmp": "BMP",
    "ico": "ICO",
    "tiff": "TIFF",
    "tif": "TIFF",
    "gif": "GIF",
}

IMG_MIME = {
    "png": "image/png",
    "jpg": "image/jpeg",
    "jpeg": "image/jpeg",
    "webp": "image/webp",
    "bmp": "image/bmp",
    "ico": "image/x-icon",
    "tiff": "image/tiff",
    "tif": "image/tiff",
    "gif": "image/gif",
}


def image_convert(
    input_path: str,
    output_format: str,
    quality: int = 90,
    width: int | None = None,
    height: int | None = None,
) -> tuple[bytes, str]:
    img = Image.open(input_path)

    if img.mode == "RGBA" and output_format in ("jpg", "jpeg"):
        bg = Image.new("RGB", img.size, (255, 255, 255))
        bg.paste(img, mask=img.split()[3])
        img = bg
    elif img.mode not in ("RGB", "RGBA") and output_format in ("jpg", "jpeg", "bmp", "webp"):
        img = img.convert("RGB")

    if width or height:
        orig_w, orig_h = img.size
        if width and height:
            new_w, new_h = width, height
        elif width:
            ratio = width / orig_w
            new_w, new_h = width, int(orig_h * ratio)
        else:
            ratio = height / orig_h
            new_w, new_h = int(orig_w * ratio), height
        img = img.resize((new_w, new_h), Image.LANCZOS)

    save_kwargs = {}

    if output_format in ("jpg", "jpeg"):
        save_kwargs["quality"] = quality
        save_kwargs["optimize"] = True
    elif output_format == "webp":
        save_kwargs["quality"] = quality
    elif output_format == "png":
        save_kwargs["optimize"] = True
    elif output_format == "ico":
        sizes = [(256, 256), (128, 128), (64, 64), (48, 48), (32, 32), (16, 16)]
        ico_path = input_path + ".ico"
        img.save(ico_path, format="ICO", sizes=sizes)
        with open(ico_path, "rb") as f:
            data = f.read()
        os.unlink(ico_path)
        return data, IMG_MIME[output_format]

    img_io = io.BytesIO()
    img.save(img_io, format=IMG_FORMATS[output_format], **save_kwargs)
    return img_io.getvalue(), IMG_MIME[output_format]


# ─── Audio Conversion ────────────────────────────────────────────────────────

import shutil as _shutil

def _find_ffmpeg():
    ffmpeg_path = _shutil.which("ffmpeg")
    if ffmpeg_path:
        return ffmpeg_path
    try:
        import imageio_ffmpeg
        return imageio_ffmpeg.get_ffmpeg_exe()
    except ImportError:
        pass
    raise FileNotFoundError("ffmpeg not found. Install ffmpeg or pip install imageio[ffmpeg]")

try:
    os.environ["FFMPEG_BINARY"] = _find_ffmpeg()
except FileNotFoundError:
    pass

from pydub import AudioSegment
try:
    AudioSegment.converter = _find_ffmpeg()
except FileNotFoundError:
    pass

AUDIO_FORMATS = {
    "mp3": "mp3",
    "wav": "wav",
    "aac": "aac",
    "flac": "flac",
    "ogg": "ogg",
    "m4a": "mp4",
    "wma": "mp3",
}

AUDIO_MIME = {
    "mp3": "audio/mpeg",
    "wav": "audio/wav",
    "aac": "audio/aac",
    "flac": "audio/flac",
    "ogg": "audio/ogg",
    "m4a": "audio/mp4",
    "wma": "audio/x-ms-wma",
}


def audio_convert(
    input_path: str,
    output_format: str,
    quality: int = 90,
    bitrate: str = "192k",
) -> tuple[bytes, str]:
    audio = AudioSegment.from_file(input_path)
    buf = io.BytesIO()

    if output_format == "mp3":
        audio.export(buf, format="mp3", bitrate=bitrate)
    elif output_format == "wav":
        audio.export(buf, format="wav")
    elif output_format == "aac":
        audio.export(buf, format="aac")
    elif output_format == "flac":
        audio.export(buf, format="flac")
    elif output_format == "ogg":
        audio.export(buf, format="ogg")
    elif output_format == "m4a":
        audio.export(buf, format="mp4")
    elif output_format == "wma":
        audio.export(buf, format="mp3", bitrate=bitrate)

    return buf.getvalue(), AUDIO_MIME[output_format]


# ─── Video to GIF Conversion ─────────────────────────────────────────────────

import moviepy
import moviepy.config
try:
    moviepy.config.FFMPEG_BINARY = os.environ.get("FFMPEG_BINARY") or _find_ffmpeg()
except FileNotFoundError:
    pass
from moviepy import VideoFileClip


def video_to_gif(
    input_path: str,
    fps: int = 10,
    width: int | None = None,
    start: float = 0,
    duration: float | None = None,
) -> tuple[bytes, str]:
    clip = VideoFileClip(input_path)

    if duration:
        clip = clip.subclipped(start, start + duration)
    elif start > 0:
        clip = clip.subclipped(start)

    if width:
        clip = clip.resized(width=width)
    elif clip.w > 480:
        clip = clip.resized(width=480)

    gif_path = input_path + ".gif"
    clip.write_gif(gif_path, fps=fps, optimization=True)
    clip.close()

    with open(gif_path, "rb") as f:
        data = f.read()
    os.unlink(gif_path)

    return data, "image/gif"


# ─── Spreadsheet Conversion ──────────────────────────────────────────────────

import pandas as pd
from openpyxl import Workbook


def spreadsheet_convert(
    input_path: str,
    output_format: str,
) -> tuple[bytes, str, str]:
    ext = Path(input_path).suffix.lower()

    if ext in (".xlsx", ".xls"):
        df = pd.read_excel(input_path)
    elif ext == ".csv":
        df = pd.read_csv(input_path)
    elif ext == ".tsv":
        df = pd.read_csv(input_path, sep="\t")
    elif ext == ".json":
        df = pd.read_json(input_path)
    else:
        raise ValueError(f"Unsupported input format: {ext}")

    if output_format == "csv":
        data = df.to_csv(index=False).encode("utf-8")
        return data, "text/csv", ".csv"
    elif output_format == "tsv":
        data = df.to_csv(index=False, sep="\t").encode("utf-8")
        return data, "text/tab-separated-values", ".tsv"
    elif output_format == "json":
        data = df.to_json(orient="records", indent=2).encode("utf-8")
        return data, "application/json", ".json"
    elif output_format == "xlsx":
        buf = io.BytesIO()
        df.to_excel(buf, index=False, engine="openpyxl")
        return buf.getvalue(), "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", ".xlsx"
    elif output_format == "html":
        data = df.to_html(index=False).encode("utf-8")
        return data, "text/html", ".html"
    elif output_format == "markdown":
        data = df.to_markdown(index=False).encode("utf-8")
        return data, "text/markdown", ".md"
    else:
        raise ValueError(f"Unsupported output format: {output_format}")


# ─── Document Conversion ─────────────────────────────────────────────────────

def document_convert(input_path: str, output_format: str) -> tuple[bytes, str, str]:
    ext = Path(input_path).suffix.lower()

    if ext == ".html" or ext == ".htm":
        content = Path(input_path).read_text(encoding="utf-8")
        from xhtml2pdf import pisa
        buf = io.BytesIO()
        pisa.CreatePDF(content, dest=buf)
        return buf.getvalue(), "application/pdf", ".pdf"
    elif ext == ".md":
        content = Path(input_path).read_text(encoding="utf-8")
        import markdown
        html = markdown.markdown(content, extensions=["tables", "fenced_code"])
        full_html = f"<html><body><style>body{{font-family:sans-serif;margin:40px;}}</style>{html}</body></html>"
        from xhtml2pdf import pisa
        buf = io.BytesIO()
        pisa.CreatePDF(full_html, dest=buf)
        return buf.getvalue(), "application/pdf", ".pdf"
    elif ext == ".docx":
        doc = Document(input_path)
        html_parts = ["<html><body><style>body{font-family:sans-serif;margin:40px;} table{border-collapse:collapse;} td,th{border:1px solid #ccc;padding:6px;}</style>"]
        for para in doc.paragraphs:
            if para.style.name.startswith("Heading"):
                level = para.style.name.replace("Heading ", "")
                html_parts.append(f"<h{level}>{para.text}</h{level}>")
            elif para.text.strip():
                html_parts.append(f"<p>{para.text}</p>")
        for table in doc.tables:
            html_parts.append("<table>")
            for i, row in enumerate(table.rows):
                tag = "th" if i == 0 else "td"
                html_parts.append("<tr>" + "".join(f"<{tag}>{c.text}</{tag}>" for c in row) + "</tr>")
            html_parts.append("</table>")
        html_parts.append("</body></html>")
        html_content = "\n".join(html_parts)

        if output_format == "html":
            return html_content.encode("utf-8"), "text/html", ".html"
        elif output_format == "pdf":
            from xhtml2pdf import pisa
            buf = io.BytesIO()
            pisa.CreatePDF(html_content, dest=buf)
            return buf.getvalue(), "application/pdf", ".pdf"
    elif ext == ".txt":
        content = Path(input_path).read_text(encoding="utf-8")
        if output_format == "html":
            html = f"<html><body><style>body{{font-family:sans-serif;margin:40px;white-space:pre-wrap;}}</style><pre>{content}</pre></body></html>"
            return html.encode("utf-8"), "text/html", ".html"
        elif output_format == "pdf":
            from xhtml2pdf import pisa
            html = f"<html><body><style>body{{font-family:sans-serif;margin:40px;white-space:pre-wrap;}}</style><pre>{content}</pre></body></html>"
            buf = io.BytesIO()
            pisa.CreatePDF(html, dest=buf)
            return buf.getvalue(), "application/pdf", ".pdf"

    raise ValueError(f"Unsupported conversion: {ext} -> {output_format}")
